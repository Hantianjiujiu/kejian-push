#!/usr/bin/env python3
"""生成《2026年6月：AI带来的国内职场真实痛点、关注焦点与对策》Word文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 标题 ──
title = doc.add_heading('2026年6月：AI带来的国内职场真实痛点、关注焦点与对策', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据来源：BBC News / 36氪 / 澎湃 / 工人日报 / 智联招聘 / 世界经济论坛 / 脉脉 / 知乎等公开发表报道与报告')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')

# ==================== 第一部分 ====================
doc.add_heading('一、五大最痛的真实痛点', level=1)

# ── 痛点 1 ──
doc.add_heading('痛点 1：岗位被替代——已从焦虑变成现实', level=2)

doc.add_paragraph('裁员数据：', style='List Bullet')
data_items = [
    '亚马逊裁员约1.6万人，甲骨文计划裁2-3万人；阿里、网易、京东"静默瘦身"',
    '世界经济论坛预测：到2030年全球替代9200万个岗位，同时创造1.7亿个新岗位',
    '但新增和消失之间隔着一道技能鸿沟',
]
for item in data_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('已受实质性冲击的岗位：', style='List Bullet')

table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['岗位', '冲击程度', '现状']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

data = [
    ['初级程序员/外包开发', '★★★★★', '大厂AI代码生成率超30%，不再招聘应届生'],
    ['客服', '★★★★★', 'AI闭环处理95%工单，团队从50人裁到5人'],
    ['翻译', '★★★★☆', '90%项目转为"机翻+人工校对"，收入腰斩'],
    ['基础设计/美工', '★★★★☆', '广告公司基础美工80%被替代'],
    ['金融分析', '★★★☆☆', 'AI半天完成过去3-4天的竞品分析'],
]
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph('')

# ── 痛点 2 ──
doc.add_heading('痛点 2：技能折旧速度赶不上技术迭代', level=2)
items = [
    '技能有效期从5-10年缩短至6-12个月',
    '47%的互联网岗位已明确要求AI能力',
    '"还没来得及学会就过时了"——两会代表总结的三大焦虑之一',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ── 痛点 3 ──
doc.add_heading('痛点 3："教AI替代自己"的道德困境', level=2)
items = [
    '公司要求员工把自己的经验萃取出来训练AI Agent',
    '员工用40分钟教AI，相当于"自己把自己开掉"',
    '已有"反蒸馏Skill"工具出现，帮员工在交差时洗掉核心知识',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ── 痛点 4 ──
doc.add_heading('痛点 4：AI疲惫症——效率高了，人更累了', level=2)
items = [
    'Token消耗量被纳入绩效考核——成"第四薪酬"',
    '专注效率降至60%（三年新低），面临倦怠风险的员工升至23%',
    '《2026年职场心理健康报告》：超70%上班族经常感到焦虑，35%因此影响睡眠',
    '89.2%的人感觉过度依赖AI后，独立思考能力下降了',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ── 痛点 5 ──
doc.add_heading('痛点 5："一人公司"梦想与007现实', level=2)
items = [
    'AI让"超级个体"成为可能，但本质是：CEO、产品、运营、客服全部自己扛，从996变007',
    '62%的Z世代职场新人没有"职场好友"，70%不向同事倾诉——情感孤岛化',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ==================== 第二部分 ====================
doc.add_heading('二、大家正在关注什么', level=1)

# ── 关注 1 ──
doc.add_heading('关注 1：AI到底杀死岗位还是创造岗位？', level=2)
doc.add_paragraph('真实数据（而非贩卖焦虑）：', style='List Bullet')

table2 = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['指标', '数据']):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
data2 = [
    ['真正担心被AI替代的人', '仅4.7%（CTR调研）'],
    ['正在主动学习AI的人', '70.11%'],
    ['企业缩编比例', '仅14.7%'],
    ['企业因AI扩张而增加招聘', '18.3%'],
    ['AI产品经理岗位增长', '同比↑81%'],
]
for i, row_data in enumerate(data2):
    for j, val in enumerate(row_data):
        table2.rows[i+1].cells[j].text = val

doc.add_paragraph('')

# ── 关注 2 ──
doc.add_heading('关注 2：35岁不是危机，是机会', level=2)
doc.add_paragraph(
    '脉脉创始人林凡判断：35+人群正在价值重估。AI可以代劳繁琐的执行，但35+沉淀的行业经验'
    '（Skill/判断力/人脉把控）恰恰是AI无法替代的核心资产。39.13%的35+员工的高阶AI用法率反而高于年轻人。'
)

# ── 关注 3 ──
doc.add_heading('关注 3：杭州判例——AI不能作为裁员理由', level=2)
doc.add_paragraph(
    '杭州法院裁定：企业引入AI工具后解雇员工属违法行为。"技术升级的风险应由企业承担，'
    '不能转嫁给劳动者"——这是2026年最重要的劳动法标杆。'
)

# ── 关注 4 ──
doc.add_heading('关注 4：从"怕AI抢饭碗"到"怕AI替你做决定"', level=2)
doc.add_paragraph(
    'IDC Q3报告：35%中大型企业已让AI智能体"自主决策上岗"（两年前仅8%）。'
    '新的恐惧是：当老板/管理者习惯"拍板前先问AI"，人类的自主判断力是否在退化？'
)

# ==================== 第三部分 ====================
doc.add_heading('三、务实可行的解决方案', level=1)

doc.add_heading('对个人（立即可做）', level=2)

doc.add_heading('第一步：停止空焦虑，开始用AI', level=3)
doc.add_paragraph('"焦虑的大多是听说过、没真用过的人"——真正用起来的人反而发现了新的可能性。')
doc.add_paragraph('具体动作：')
items = [
    '每天花30分钟把AI当工具用（写周报、做表格、总结文档）',
    '重点不是"学AI"，而是用AI放大你已有的专业能力',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('第二步：打造"AI + 行业经验"的复合护城河', level=3)
doc.add_paragraph('企业最需要的三项能力（招聘数据）：')
items = [
    '实践应用与工程化能力（38.7%）',
    '专业判断与决策力（21%）',
    '创新与解决问题能力（17%）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph('不是"学Python"，而是在你自己的领域里，比别人更会调用AI。')

doc.add_heading('第三步：保护独立思考', level=3)
doc.add_paragraph(
    '把AI的初次回答当作思考的起点而非终点。省下的时间应该投入到需要人类判断力的环节，'
    '而不是无脑接受AI输出。'
)

doc.add_heading('你的课程资产如何切入这些痛点', level=2)

doc.add_paragraph(
    '你拥有35节《实务领导力》课程 + 35条宣传视频 + 课程展示网站 + GitHub开源仓库，'
    '正好切中2026年最核心的四个需求：'
)

table3 = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['痛点', '你的课程的解决方案']):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
data3 = [
    ['技能折旧焦虑', '领导力是"无法被AI替代的能力"——这正是你代总序的核心论点'],
    ['35+价值重估', '课程体系（认识自我→改变自我→领导他人→团队→升华）直接帮人完成"从执行者到管理者"的跃迁'],
    ['超级个体困境', '刘邦用人之道、6种权力、团队管理——教人从"自己扛"到"带团队"'],
    ['职业天花板', '人生趋势图、超预期管理、副部级后台——帮人看清47岁天花板并提前布局'],
]
for i, row_data in enumerate(data3):
    for j, val in enumerate(row_data):
        table3.rows[i+1].cells[j].text = val

doc.add_paragraph('')

doc.add_heading('建议的下一步传播策略', level=2)
items = [
    '把35条promo视频分发到抖音/视频号/B站/小红书，每条标题已经按爆款优化',
    'GitHub README已带微信引流，每个平台发视频时在评论区置顶GitHub链接',
    '用你刚安装的GordenSuperPPTSkill，把每节课内容快速生成图片式PPT → 发小红书/即刻（图文笔记）',
    '打"反焦虑"标签：不要贩卖焦虑，你的定位是"AI时代，领导力是你的护城河"——和市面上99%的焦虑营销彻底拉开差距',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ==================== 页脚 / 来源 ====================
doc.add_page_break()
doc.add_heading('数据来源', level=1)
sources = [
    ('BBC News: 当AI智能体开始上岗——2026年春天里，三个中国人的兴奋与恐惧',
     'https://www.bbc.com/zhongwen/articles/crm8mmg0y3eo/simp'),
    ('36氪: 2026的春天，AI正在杀死普通职场人',
     'https://36kr.com/p/3741012034351881'),
    ('36氪: 那些刷屏的职场焦虑，该醒醒了',
     'https://36kr.com/p/3838298342197766'),
    ('澎湃: 按要求「蒸馏」自己后，他们被裁了',
     'https://www.thepaper.cn/newsDetail_forward_33162759'),
    ('工人日报: AI焦虑，是被替代还是新同事',
     'https://www.workercn.cn/c/2026-03-07/8749641.shtml'),
    ('中国经济网: AI浪潮下，大厂产生降本裁员冲动',
     'http://www.ce.cn/xwzx/gnsz/gdxw/202606/t20260606_3014702.shtml'),
    ('智联招聘: AI就业冰火交替，主动拥抱才是出路',
     'https://tech.china.com.cn/sx/20260522/412894.shtml'),
    ('红网: AI到底抢走了谁的岗位',
     'https://hlj.rednet.cn/content/646047/61/15799082.html'),
    ('澎湃: 第一批用AI的人，已经染上了AI疲惫症',
     'https://www.thepaper.cn/newsDetail_forward_32916550'),
    ('Sputnik/BBC: 中国开始保护就业岗位免受人工智能冲击',
     'https://sputniknews.cn/20260505/1071097691.html'),
    ('界面新闻: 有了AI，人们却过得越来越累了？',
     'https://m.jiemian.com/article/14270730.html'),
    ('香港01: AI大军压境抢饭碗，打工仔技荒心更慌',
     'https://global.hk01.com'),
]

# 添加分割线
for title_text, url in sources:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title_text}')
    run.font.size = Pt(9)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'  {url}')
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x33, 0x66, 0x99)

# ── 保存 ──
out_path = r'F:\kejian\AI职场痛点与对策-2026年6月.docx'
doc.save(out_path)
print(f'已保存到: {out_path}')
